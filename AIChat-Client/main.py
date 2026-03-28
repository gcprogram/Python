import tkinter as tk
from tkinter import scrolledtext, filedialog, messagebox, ttk
import time
import requests
import json
import os
import threading
import logging
import re
from docx import Document # Benötigt
from fpdf import FPDF  # Benötigt pip install fpdf
from DocumentProcessor import DocumentProcessor # Verarbeitung von Dokumenten für Übersetzungen

logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s | %(levelname)s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)
log = logging.getLogger(__name__)

CONFIG_FILE = "config.json"

LANG_MAP = {
    "en": "English", "de": "German", "cs": "Czech", "es": "Spanish",
    "fr": "French", "it": "Italian", "pt": "Portuguese", "ru": "Russian",
    "zh": "Chinese", "ja": "Japanese", "ko": "Korean", "tr": "Turkish"
}
# Inverse Map für die GUI-Auswahl
UI_LANG_LIST = sorted([f"{name} ({code})" for code, name in LANG_MAP.items()])

LANG_MAP_FULL = {
    "af": "Afrikaans", "ar": "Arabic", "az": "Azerbaijani", "be": "Belarusian",
    "bg": "Bulgarian", "bn": "Bengali", "ca": "Catalan", "cs": "Czech",
    "cy": "Welsh", "da": "Danish", "de": "German", "el": "Greek",
    "en": "English", "es": "Spanish", "et": "Estonian", "fa": "Persian",
    "fi": "Finnish", "fr": "French", "ga": "Irish", "gl": "Galician",
    "gu": "Gujarati", "he": "Hebrew", "hi": "Hindi", "hr": "Croatian",
    "hu": "Hungarian", "hy": "Armenian", "id": "Indonesian", "is": "Icelandic",
    "it": "Italian", "ja": "Japanese", "ka": "Georgian", "kk": "Kazakh",
    "km": "Khmer", "kn": "Kannada", "ko": "Korean", "lt": "Lithuanian",
    "lv": "Latvian", "mk": "Macedonian", "ml": "Malayalam", "mn": "Mongolian",
    "mr": "Marathi", "ms": "Malay", "mt": "Maltese", "my": "Burmese",
    "nl": "Dutch", "no": "Norwegian", "pa": "Punjabi", "pl": "Polish",
    "pt": "Portuguese", "ro": "Romanian", "ru": "Russian", "sk": "Slovak",
    "sl": "Slovenian", "sq": "Albanian", "sr": "Serbian", "sv": "Swedish",
    "sw": "Swahili", "ta": "Tamil", "te": "Telugu", "th": "Thai",
    "tr": "Turkish", "uk": "Ukrainian", "ur": "Urdu", "uz": "Uzbek",
    "vi": "Vietnamese", "zh": "Chinese"
}

class ChatApp:
    #
    # Initialisierung und Hauptfenster
    #
    def __init__(self, master):
        self.master = master
        master.title("KI Chat und Übersetzungs-Client")

        # Einstellungen laden
        self.load_config()
        self.processor = DocumentProcessor(max_chunk_chars=self.max_chunk_chars)

        # UI Komponenten
        self.chat_area = scrolledtext.ScrolledText(master, wrap=tk.WORD, state='normal')
        self.chat_area.grid(row=0, column=0, columnspan=4, padx=10, sticky="we")

        #self.user_input = tk.Entry(master)
        self.user_input = tk.Text(master, height=3, width=40)
        self.user_input.grid(row=1, column=1, columnspan=3, rowspan=3, padx=10, pady=5,sticky="we")

        self.settings_button = tk.Button(master, text="Einstellungen", command=self.open_settings)
        self.settings_button.grid(row=22, column=0)

        self.file_button = tk.Button(master, text="Datei übersetzen (Pipeline)", command=self.upload_file)
        self.file_button.grid(row=22, column=1)

        self.file_button = tk.Button(master, text="Chat speichern", command=self.save_chat)
        self.file_button.grid(row=22, column=2)

        self.send_button = tk.Button(master, text="Senden", command=self.send_message)
        self.send_button.grid(row=22, column=3)

        # Status Bar & Progress
        self.setup_statusbar()

        # Formatierung
        self.chat_area.tag_configure("reasoning", background="#f5f5f5", foreground="#666666")
        self.chat_area.tag_configure("message", background="#e0e0e0", foreground="black")
        self.chat_area.tag_configure("bold", font=("Arial", 10, "bold"))
        self.chat_area.tag_configure("system", foreground="blue", font=("Arial", 10, "italic"))

        self.last_type = None
        self.load_start_time = None
        self.response_id = ""

    #
    # Bringt die Status Bar unter das Hauptfenster
    #
    def setup_statusbar(self):
        self.style = ttk.Style()
        self.style.theme_use('default')
        self.style.configure("green.Horizontal.TProgressbar", foreground='green', background='green')
        self.style.configure("red.Horizontal.TProgressbar", foreground='red', background='red')

        self.status_frame = tk.Frame(self.master, bd=1, relief=tk.SUNKEN)
        self.status_frame.grid(row=23, column=0, columnspan=4, padx=10, pady=10, sticky="we")

        self.status_var = tk.StringVar(value="Bereit")
        self.status_label = tk.Label(self.status_frame, textvariable=self.status_var, anchor=tk.W)
        self.status_label.grid(row=0, column=0, columnspan=2, padx=5)

        self.progress = ttk.Progressbar(self.status_frame, orient=tk.HORIZONTAL, length=150, mode='determinate',
                                        style="green.Horizontal.TProgressbar")
        self.progress.grid(row=0, column=2, padx=5, pady=2, sticky="e")

    #
    # Konfiguration laden
    #
    def load_config(self, filename=CONFIG_FILE):
        if os.path.exists(filename):
            try:
                with open(filename, "r", encoding="utf-8") as f:
                    config = json.load(f)
                    self.api_key = config.get("api_key", "")
                    self.ip = config.get("ip", "")
                    self.port = config.get("port", "")
                    self.model = config.get("model", "")
                    self.models = config.get("models", [])
                    self.translation_prompt = config.get("translation_prompt",
                                                         "Übersetze den folgenden Text präzise ins Deutsche. Behalte Markdown-Formatierungen bei.")
                    self.max_chunk_chars = config.get("max_chunk_chars", 3500)
                    self.lang_source = config.get("lang_source", "en")
                    self.lang_target = config.get("lang_target", "de-DE")

                # Falls wir eine externe Datei geladen haben, aktualisieren wir ggf. die UI-Felder,
                # sofern das Einstellungsfenster offen ist.
                log.info(f"Konfiguration aus {filename} geladen.")
            except Exception as e:
                log.error(f"Fehler beim Laden der Config {filename}: {e}")
        else:
            # Default Werte falls Datei nicht existiert
            self.api_key = self.ip = self.port = self.model = ""
            self.models = []
            self.translation_prompt = "Übersetze den folgenden Text präzise ins Deutsche."
            self.max_chunk_chars = 3500
            self.lang_source = "en"
            self.lang_target = "de-DE"
    #
    # Konfiguration speichern
    #
    def save_config(self):
        # Zielpfad abfragen
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON Konfiguration", "*.json")],
            initialfile="config.json",
            title="Konfiguration speichern unter..."
        )

        if not file_path:
            return

        # Daten aus UI extrahieren
        selection_s = self.source_lang_combo.get()
        self.lang_source = selection_s[selection_s.find("(") + 1: selection_s.find(")")]
        selection_t = self.target_lang_combo.get()
        self.lang_target = selection_t[selection_t.find("(") + 1: selection_t.find(")")]

        config = {
            "api_key": self.token_entry.get(),
            "ip": self.ip_entry.get(),
            "port": self.port_entry.get(),
            "model": self.model_var.get(),
            "models": self.models,
            "lang_source": self.lang_source,
            "lang_target": self.lang_target,
            "translation_prompt": self.prompt_text.get("1.0", tk.END).strip(),
            "max_chunk_chars": int(self.chunk_entry.get())
        }

        # 1. In gewählte Datei speichern
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=4)

        # 2. Immer auch in die Standard-Config kopieren (für Auto-Load beim Start)
        if file_path != CONFIG_FILE:
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(config, f, indent=4)

        self.load_config(file_path)  # Variablen im Objekt aktualisieren
        self.processor.max_chunk_chars = self.max_chunk_chars
        messagebox.showinfo("Erfolg", f"Einstellungen in {os.path.basename(file_path)} gespeichert.")

    def manual_load_config(self, settings_window):
        file_path = filedialog.askopenfilename(
            filetypes=[("JSON Konfiguration", "*.json")],
            title="Konfiguration laden"
        )
        if file_path:
            self.load_config(file_path)
            # Schließe das Einstellungsfenster und öffne es neu, um die Felder zu aktualisieren
            settings_window.destroy()
            self.open_settings()

    #
    # Einstellungsfenster bauen
    #
    def open_settings(self):
        settings_window = tk.Toplevel(self.master)
        settings_window.title("Einstellungen")
        #settings_window.geometry("600x500")

        # IP-Adresse
        tk.Label(settings_window, text="IP-Adresse:").grid(row=0, column=0, padx=10, pady=5, sticky='w')
        self.ip_entry = tk.Entry(settings_window)
        self.ip_entry.insert(0, self.ip)
        self.ip_entry.grid(row=0, column=1, padx=10, sticky='w')

        # Port
        tk.Label(settings_window, text="Port:").grid(row=1, column=0, padx=10, pady=5, sticky='w')
        self.port_entry = tk.Entry(settings_window)
        self.port_entry.insert(0, self.port)
        self.port_entry.grid(row=1, column=1, padx=10, sticky='w')

        # API-Schlüssel
        tk.Label(settings_window, text="API-Schlüssel:").grid(row=2, column=0, padx=10, pady=5, sticky='w')
        self.token_entry = tk.Entry(settings_window, show="*")
        self.token_entry.insert(0, self.api_key)
        self.token_entry.grid(row=2, column=1, padx=10, sticky='w')

        # Übersetzungs-Prompt
        tk.Label(settings_window, text="Übersetzung Prompt:").grid(row=3, column=0, padx=10, pady=5, sticky='nw')
        self.prompt_text = tk.Text(settings_window, height=4, width=41)
        self.prompt_text.insert("1.0", self.translation_prompt)
        self.prompt_text.grid(row=3, column=1, padx=10, pady=5, sticky='w')

        # Chunk Größe
        tk.Label(settings_window, text="Teilstücklänge:").grid(row=4, column=0, padx=10, pady=5, sticky='w')
        self.chunk_entry = tk.Entry(settings_window)
        self.chunk_entry.insert(0, str(self.max_chunk_chars))
        self.chunk_entry.grid(row=4, column=1, padx=10, sticky='w')

        tk.Label(settings_window, text="Original/Zielsprache:").grid(row=5, column=0, padx=10, pady=5)
#        options_lang = [ "de-DE", "en", "es-ES", "cs-CZ", "fr", "it", "nl", "pl", "ru", "zh-CN",
#                         "ja-JP","ko-KR", "tr-TR", "pl-PL", "sv-SE" ]
#        self.langsource_var = tk.StringVar(settings_window, value=self.lang_source)
#        self.langtarget_var = tk.StringVar(settings_window, value=self.lang_target)
#        self.menu_langsource = tk.OptionMenu(settings_window, self.langsource_var, *options_lang)
#        self.menu_langtarget = tk.OptionMenu(settings_window, self.langtarget_var, *options_lang)
#        self.menu_langsource.grid(row=5, column=1, padx=10, pady=5, sticky="w")
#        self.menu_langtarget.grid(row=5, column=1, padx=10, pady=5, sticky="e")

        # Neue Combobox für Sprachenwahl
        self.source_lang_combo = ttk.Combobox(settings_window, values=UI_LANG_LIST)
        self.target_lang_combo = ttk.Combobox(settings_window, values=UI_LANG_LIST)
        # Setze aktuellen Wert (muss aus Code wieder in den UI-Namen gewandelt werden)
        current_source_val = f"{LANG_MAP.get(self.lang_source)} ({self.lang_source})"
        self.source_lang_combo.set(current_source_val)
        current_target_val = f"{LANG_MAP.get(self.lang_target)} ({self.lang_target})"
        self.source_lang_combo.set(current_source_val)
        self.target_lang_combo.set(current_target_val)
        self.source_lang_combo.grid(row=5, column=1, sticky="w", padx=10)
        self.target_lang_combo.grid(row=5, column=1, sticky="e", padx=10)



        # Modell Auswahl
        tk.Label(settings_window, text="Modell:").grid(row=6, column=0, padx=10, pady=5)
        self.model_var = tk.StringVar(settings_window, value=self.model)
        menu_options = self.models if self.models else ["Keine Modelle"]
        self.model_menu = tk.OptionMenu(settings_window, self.model_var, *menu_options)
        self.model_menu.grid(row=6, column=1, padx=10, sticky='w')

        # Container-Frame für die Buttons ganz unten
        button_frame = tk.Frame(settings_window)
        button_frame.grid(row=7, column=0, columnspan=2, pady=20)

        # Modelle laden Button
        self.load_models_button = tk.Button(button_frame, text="Modelle laden", command=self.load_models)
        self.load_models_button.pack(side=tk.LEFT, padx=5)

        # NEU: Konfiguration Laden Button
        load_btn = tk.Button(
            button_frame,
            text="Config Laden",
            command=lambda: self.manual_load_config(settings_window),
            width=12
        )

        load_btn.pack(side=tk.LEFT, padx=5)

        # Speichern Button
        save_btn = tk.Button(
            button_frame,
            text="Speichern",
            command=lambda: [self.save_config(), settings_window.destroy()],
            width=12,
            bg = "#e1f5fe"
        )
        save_btn.pack(side=tk.LEFT, padx=5)

        # Schließen / Abbrechen Button
        close_btn = tk.Button(
            button_frame,
            text="Schließen",
            command=settings_window.destroy,
            width=12
        )
        close_btn.pack(side=tk.RIGHT, padx=5)


    #
    # Holt KI-Model vom KI-Server und packt sie in die Dropdownliste.
    #
    def load_models(self):
        log.info("> load_models()")
        try:
            # 1. Daten abrufen (Beispielhaft via Requests)
            # Ersetze die URL durch deine tatsächliche API-Endpunkt-URL
            # response = requests.get(f"http://{self.ip}:{self.port}/v1/models")
            # data = response.json()
            data = self._get_models()
            log.info(f"Modelle stehen zur Verfügung: {data}")

            # 2. Die "key"-Elemente extrahieren
            log.info(f"{data}")
            # Nur hinzufügen, wenn der Key "key" auch wirklich existiert
            self.models = [model["key"] for model in data if "key" in model]
            log.info(f"keys: {self.models}")

            # 3. Das OptionMenu-Widget aktualisieren
            menu = self.model_menu["menu"]
            menu.delete(0, "end")  # Alte Einträge entfernen

            for model_key in self.models:
                # Hier fügen wir jeden Key hinzu und sorgen dafür,
                # dass beim Klick die model_var gesetzt wird
                menu.add_command(
                    label=model_key,
                    command=lambda value=model_key: self.model_var.set(value)
                )

            # 4. Optional: Den ersten Eintrag automatisch auswählen, falls nichts selektiert ist
            if self.models and not self.model_var.get():
                self.model_var.set(self.models[0])

            log.info(f"{len(self.models)} Modelle erfolgreich geladen.")
            log.info("< load_models()")

        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Laden der Modellliste: {e}")
            log.error("< load_models()")

    #
    # Requested die Modelle vom Server.
    #
    def _get_models(self):
        log.info("> get_models()")
        url = f"http://{self.ip}:{self.port}/api/v1/models"
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        try:
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code == 200:
                log.info("< get_models()")
                self.master.after(0, lambda: [
                    self.status_var.set("Erfolgreiche Modellabfrage."),
                    self.progress.configure(style="green.Horizontal.TProgressbar", value=100)
                ])
                return response.json().get("models", [])
            elif response.status_code == 401:
                log.error("< get_models(): Falsches API Token.")
                self.master.after(0, lambda: [
                    self.status_var.set("Nicht authorisiert."),
                    self.progress.configure(style="red.Horizontal.TProgressbar", value=100)
                ])
            else:
                log.error("< get_models(): Falscher Statuswert bei Laden der Modellliste.")
                self.master.after(0, lambda: [
                    self.status_var.set("Falscher Statuswert bei Laden der Modellliste"),
                    self.progress.configure(style="red.Horizontal.TProgressbar", value=100)
                ])
                return []

        except requests.exceptions.ConnectTimeout:
            log.error("Connection Timeout.")

        except requests.exceptions.ReadTimeout:
            log.error("Der Server hat die Verbindung akzeptiert, aber braucht zu lange zum Antworten.")

        except requests.exceptions.Timeout:
            # Dies fängt beide obigen Fälle ab
            log.error("Allgemeiner Zeitüberschreitungsfehler.")

        except requests.exceptions.RequestException as e:
            # Fängt alle anderen Requests-Probleme ab (z.B. falsche URL)
            log.error(f"Ein schwerwiegender Fehler ist aufgetreten:", e)
        except Exception as e:
                err_short = str(e)[:40] + "..." if len(str(e)) > 40 else str(e)
                self.master.after(0, lambda: [
                    self.status_var.set(f"< get_models(): {err_short}"),
                    self.progress.configure(style="red.Horizontal.TProgressbar", value=100)
                ])

    #
    # Zentrale Methode zum Aufbau der Payload für die Modell-Anfrage
    #
    def _prepare_payload(self, message, is_translation=False):
        """Entscheidet, welches Format (Standard oder TranslateGemma) gesendet wird."""

        # Check: Ist es ein TranslateGemma Modell?
        if "translategemma" in self.model.lower():
            # Das spezialisierte Template
            structured_content = [
                {
                    "type": "text",
                    "source_lang_code": self.lang_source,
                    "target_lang_code": self.lang_target,
                    "text": message
                }
            ]
            log.info(f"translategemma: content={structured_content}")
            # In den meisten APIs wird dies als 'input' oder Teil der 'messages' gesendet
            return {
                "model": self.model,
                "input": structured_content,  # Viele Backends akzeptieren hier das Objekt
                "stream": not is_translation  # Sync bei Übersetzung, Stream beim Chat
            }
        else:
            # Standard-Format für normale LLMs
            return {
                "model": self.model,
                "input": message,
                "stream": not is_translation
            }


    def chat_with_ai(self, message):
        log.info(f"> chat_with_ai() gestartet für Modell: {self.model}")
        url = f"http://{self.ip}:{self.port}/api/v1/chat"
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }

        # Sicherstellen, dass die Payload so aussieht, wie die API sie braucht.
        # Falls es ein OpenAI-kompatibler Endpunkt ist, müsste es 'messages' sein.
        payload = {
            "input": message,  # Beibehalten, falls deine API das so braucht
            "model": self.model,
            "stream": True
        }

        if self.response_id:
            payload["previous_response_id"] = self.response_id

        try:
            # 1. Timeout hinzugefügt (10 Sek auf Verbindung, None auf Stream)
            log.info(f"Sende Request an {url}...")
            response = requests.post(url, headers=headers, json=payload, stream=True, timeout=(10, None))

            # 2. Sofort den Status loggen!
            log.info(f"Response erhalten: Status {response.status_code}")

            if response.status_code == 200:
                for line in response.iter_lines():
                    if line:
                        decoded_line = line.decode('utf-8').strip()
                        if decoded_line.startswith("data: "):
                            json_str = decoded_line[6:]
                            # Innerhalb der for-Schleife in chat_with_ai:
                            try:
                                chunk = json.loads(json_str)
                                chunk_type = chunk.get("type", "")

                                # FALL 1: Modell lädt noch (Progress-Anzeige)
                                if chunk_type == "model_load.progress":
                                    p_value = chunk.get("progress", 0)

                                    if self.load_start_time is None:
                                        self.load_start_time = time.time()

                                    # Berechnung der Restzeit
                                    elapsed = time.time() - self.load_start_time
                                    if p_value > 0:
                                        total_estimated_time = elapsed / p_value
                                        remaining_time = max(0, total_estimated_time - elapsed)
                                        time_str = f"{int(remaining_time)}s verbleibend"
                                    else:
                                        time_str = "Lade..."

                                    def update_p(v=p_value, t=time_str):
                                        self.progress["value"] = v * 100
                                        self.progress["style"] = "green.Horizontal.TProgressbar"
                                        self.status_var.set(f"Modell lädt... ({t})")

                                    self.master.after(0, update_p)

                                # FALL 2: Laden beendet / Prompt-Verarbeitung
                                elif chunk_type == "model_load.end":
                                    self.load_start_time = None  # Reset für nächstes Mal
                                    self.master.after(0, lambda: [self.status_var.set("Modell bereit."),
                                                                  self.progress.configure(value=100)])


                                # FALL 3: Echtes Streaming (Reasoning / Message)
                                # Delta-Events für Reasoning und Message (wie besprochen)
                                elif ".delta" in chunk_type:
                                    # Status-Bar zurücksetzen, wenn die erste Antwort kommt
                                    if self.status_var.get() != "KI schreibt...":
                                        self.master.after(0, lambda: self.status_var.set("KI schreibt..."))
                                    item = {
                                        "type": chunk_type.split(".")[0],
                                        "content": chunk.get("content", "")
                                    }
                                    self.master.after(0, self._update_chat_ui, item)

                                # Ende des Chats
                                elif chunk_type == "chat.end":
                                    self.master.after(0, lambda: self.status_var.set("Bereit"))

                            except json.JSONDecodeError:
                                continue
            else:
                err_msg = f"Server Fehler: {response.status_code} - {response.text}"
                log.error(err_msg)
                self.master.after(0, lambda: self.chat_area.insert(tk.END, f"\n{err_msg}\n"))

        except requests.exceptions.Timeout:
            log.error("Timeout: Der Server antwortet nicht schnell genug.")
            self.master.after(0, lambda: self.chat_area.insert(tk.END,
                                                               "\nFehler: Verbindung zum Server zeitüberschreitung.\n"))
        except Exception as e:
            err_short = str(e)[:40] + "..." if len(str(e)) > 40 else str(e)
            self.master.after(0, lambda: [
                self.status_var.set(f"Verbindungsfehler: {err_short}"),
                self.progress.configure(style="red.Horizontal.TProgressbar", value=100)
            ])

    def _update_chat_ui(self, item):
        content_type = item.get("type")
        content_text = item.get("content", "")

        if not content_text:
            return

        # Spezialbehandlung für den Ladevorgang
        if content_type == "loading":
            # Wir löschen die letzte Zeile nicht (kompliziert in Tkinter),
            # aber wir markieren es als Systemnachricht
            if self.last_type != "loading":
                self.chat_area.insert(tk.END, "\n")
                self.last_type = "loading"

            # Cursor an den Anfang der Zeile setzen ist im Text-Widget schwierig,
            # daher hängen wir es hier einfach an oder nutzen ein Label.
            # Für den Anfang reicht eine einfache Statuszeile:
            self.chat_area.insert(tk.END, content_text)
            self.chat_area.see(tk.END)
            return

        # Rest der Logik für Reasoning und Message...
        if content_type != self.last_type:
            # Zeilenumbruch vor neuem Block, außer ganz am Anfang
            if self.last_type is not None:
                self.chat_area.insert(tk.END, "\n")

            header = "Reasoning:\n" if content_type == "reasoning" else "\n"
            tag = "reasoning" if content_type == "reasoning" else "message"

            self.chat_area.insert(tk.END, header, (tag, "bold"))
            self.last_type = content_type

        # Den Text-Schnipsel mit dem richtigen Tag (Farbe) einfügen
        tag = "reasoning" if content_type == "reasoning" else "message"
        self.chat_area.insert(tk.END, content_text, tag)

        # Sofort zum Ende scrollen
        self.chat_area.see(tk.END)

    def send_message(self):
        log.info("> send_message(): sending message")
        user_message = self.user_input.get("1.0", tk.END).strip()
        if not user_message:
            return

        self.chat_area.insert(tk.END, f"\n{user_message}\n")
        self.user_input.delete("1.0", tk.END)
        self.last_type = None  # Reset für die neue Antwort

        # Thread starten, damit die GUI flüssig bleibt
        threading.Thread(target=self.chat_with_ai, args=(user_message,), daemon=True).start()
        log.info("< send_message(): after threading started")

    def upload_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Dokumente", "*.txt *.pdf *.docx *.epub")])
        if file_path:
            threading.Thread(target=self.run_translation_pipeline, args=(file_path,), daemon=True).start()

    def run_translation_pipeline(self, file_path):
        try:
            is_gemma = "translategemma" in self.model.lower()
            self.master.after(0, lambda: self.status_var.set("Extrahiere Text..."))

            # 1. Text extrahieren & Chunking
            full_text = self.processor.file_to_markdown(file_path)
            if not full_text or len(full_text.strip()) == 0:
                raise ValueError("Die Datei konnte nicht gelesen werden oder ist leer.")
            if is_gemma:
                chunks = self.processor.create_smart_chunks(full_text, overlap_chars=0)
            else:
                chunks = self.processor.create_smart_chunks(full_text, overlap_chars=300)

            if not chunks:
                raise ValueError("Das Chunking hat keine Abschnitte erzeugt.")

            final_glossary = {}
            if not is_gemma:
                # Glossar-Logik (nur für Standard-Modelle)
                self.master.after(0, lambda: self.status_var.set("KI analysiert Begriffe..."))
                ai_proposal = self.generate_dictionary(chunks[0])
                merged_dict = self.merge_dictionary_with_ai_proposal(ai_proposal)

                display_text = "".join([f"{k.capitalize()}: {v}\n" for k, v in sorted(merged_dict.items())])
                self.dictionary_ready_event = threading.Event()
                self.master.after(0, lambda: self.show_dictionary_edit_window(display_text))
                self.dictionary_ready_event.wait()

                # Glossar parsen
                for line in self.final_dict.split("\n"):
                    if ":" in line:
                        parts = line.split(":", 1)
                        final_glossary[parts[0].strip().lower()] = parts[1].strip()
                self.save_dictionary_file(final_glossary)

            # 2. Übersetzungsschleife (Flach, keine verschachtelten Loops!)
            self.master.after(0, lambda: self.progress.configure(value=0))

            for i, chunk in enumerate(chunks):
                # FALLUNTERSCHEIDUNG: Falls create_smart_chunks Tupel liefert (context, content)
                if isinstance(chunk, tuple):
                    context, content = chunk
                else:
                    content = chunk
                self.master.after(0, lambda idx=i: self.status_var.set(f"Übersetze Teil {idx + 1}/{len(chunks)}..."))

                attempts = 0
                success = False
                translation = ""

                # Prompt-Erstellung
                if is_gemma:
                    current_prompt = content  # send_sync_request baut das Template
                else:
                    current_prompt = (
                        f"SYSTEM: {self.translation_prompt}\n"
                        f"GLOSSAR: {final_glossary}\n"
                        f"TEXT:\n{content}"
                    )

                # Retry-Logik bei Abbruch
                while attempts < 2 and not success:
                    translation = self.send_sync_request(current_prompt, timeout=300)
                    # Qualitätscheck: Endet die Übersetzung abrupt?
                    valid_punc = ".!?;:»«\"ˮ"
                    source_ends_punc = content.strip()[-1] in valid_punc
                    target_ends_punc = translation.strip() and translation.strip()[-1] in valid_punc

                    if source_ends_punc and not target_ends_punc:
                        log.warning(f"Chunk {i + 1} unvollständig, Versuch {attempts + 1}...")
                        display_status = "" if success else "[⚠️ POTENZIELL UNVOLLSTÄNDIG]\n"
                        attempts += 1
                    elif not source_ends_punc and target_ends_punc:
                        log.warning(f"Chunk {i + 1} unvollständig, Versuch {attempts + 1}...")
                        display_status = "[⚠️ TEXTFLUSS UNTERBROCHEN]\n"
                        attempts += 1
                        # Die KI hat den Satz vermutlich einfach mit einem Punkt "abgehackt"
                    else:
                        display_status = ""
                        success = True

                # UI Update für diesen Chunk
                anfang = self.processor.get_boundary_sentences(content, mode='first',count=2)
                ende = self.processor.get_boundary_sentences(content, mode='last',count=2)
                output_text = f"\n\n---Original-Anfang:\n{anfang}\n--- BEGINN ABSCHNITT {i + 1}\n{translation}\n--- ENDE ABSCHNITT {i+1}\n{display_status}--- Original-Ende:\n{ende}\n"

                self.master.after(0, lambda text=output_text: self.chat_area.insert(tk.END, text))
                self.master.after(0, lambda val=(i + 1) * 100 / len(chunks): self.progress.configure(value=val))

            self.master.after(0, lambda: self.status_var.set("Übersetzung abgeschlossen."))

        except Exception as e:
            log.exception("Pipeline Fehler")
            self.master.after(0, lambda err=e: messagebox.showerror("Fehler", f"Pipeline abgebrochen: {err}"))

    def show_dictionary_edit_window(self, initial_text):
        """Öffnet ein Fenster zur manuellen Korrektur des Glossars."""
        edit_win = tk.Toplevel(self.master)
        edit_win.title("Dictionary prüfen & korrigieren")
        edit_win.geometry("400x500")

        tk.Label(edit_win, text="Bitte Dictionary für die Übersetzung korrigieren:", font=("Arial", 10, "bold")).pack(
            pady=5)

        txt_edit = tk.Text(edit_win, wrap=tk.WORD)
        txt_edit.insert("1.0", initial_text)
        txt_edit.pack(expand=True, fill='both', padx=10, pady=5)

        def on_confirm():
            self.final_dict = txt_edit.get("1.0", tk.END).strip()
            edit_win.destroy()
            self.dictionary_ready_event.set()  # Pipeline fortsetzen

        tk.Button(edit_win, text="Dictionary bestätigen & Übersetzung starten",
                  bg="green", fg="white", command=on_confirm).pack(pady=10)

        # Falls das Fenster geschlossen wird ohne Bestätigung:
        edit_win.protocol("WM_DELETE_WINDOW", on_confirm)

    def send_sync_request(self, prompt,timeout=180):
        """Sendet einen Request und erkennt automatisch das nötige Format."""
        is_gemma = "translategemma" in self.model.lower()
        # Wechsle bei Gemma auf den /v1/completions endpoint (Text-Basis)
        # Für andere Modelle bleibt es bei /v1/chat/completions
        endpoint = "/v1/completions" if is_gemma else "/v1/chat/completions"
        url = f"http://{self.ip}:{self.port}{endpoint}"
        headers = {"Authorization": f"Bearer {self.api_key}", "Content-Type": "application/json"}
        s_name = LANG_MAP.get(self.lang_source, "Original Language")
        t_name = LANG_MAP.get(self.lang_target, "Target Language")

        # Payload zusammenbauen
        if is_gemma:
            # Wir bilden das Jinja-Template in Python nach:
            # Das ist das exakte Format, das dein Jinja-Template erzeugen würde:
            full_raw_prompt = (
                f"<start_of_turn>user\n"
                f"You are a professional {s_name} ({self.lang_source}) to {t_name} ({self.lang_target}) translator. "
                f"Your goal is to accurately convey the meaning and nuances of the original {s_name} text "
                f"while adhering to {t_name} grammar, vocabulary, and cultural sensitivities.\n"
                f"Produce only the {t_name} translation, without any additional explanations or commentary, "
                f"Preserve the Markdown format in the translation. "
                f"Please translate the following Markdown text from {s_name} into {t_name}:\n\n\n"
                f"{prompt.strip()}<end_of_turn>\n"
                f"<start_of_turn>model\n"
            )

            payload = {
                "model": self.model,
                "prompt": full_raw_prompt,  # Hier direkt als 'prompt' statt 'messages'
                "stop": ["<end_of_turn>"],
                "temperature": 0.0,
                "max_tokens": 4096,
                "stream": False
            }
        else:
            # Standard-Format für normale Modelle
            messages = [{"role": "user", "content": prompt}]
            # Payload über die neue Factory holen
            payload = self._prepare_payload(prompt, is_translation=True)

        try:
            response = requests.post(url, headers=headers, json=payload, timeout=timeout)
            res_data = response.json()

            if response.status_code != 200:
                log.error(f"Payload gesendet: {json.dumps(payload)}")
                return f"Fehler {response.status_code}: {response.text}"

                # Text-Extraktion unterscheidet sich je nach Endpoint:
            if is_gemma:
                return res_data["choices"][0]["text"].strip()  # completions nutzt .text
            else:
                return res_data["choices"][0]["message"]["content"].strip()  # chat nutzt .message.content

        except Exception as e:
            log.exception(f"Schwerer Fehler: {str(e)}")
            return f"Fehler: {str(e)}"

    def generate_dictionary(self, text_sample):
        """Fragt die KI nach den wichtigsten Begriffen."""

        dictionary_prompt = (
            "Create a glossary/dictionary for translation to German. "
            "Extract only special and technical terms, proper names, and fictional terms. Only list single or connected words."
            "Answer in format: 'Original: Translation'\n"
            "If no term is found, leave it as is.\n\n"
            f"Text Example:\n{text_sample}"
        )
        return self.send_sync_request(dictionary_prompt, timeout=200)

    def merge_dictionary_with_ai_proposal(self, ai_raw_text):
        """Vergleicht KI-Vorschlag mit Datei und fügt nur Neues hinzu."""
        existing_dictionary = self.load_dictionary_file()
        new_entries_count = 0

        # KI Antwort parsen (erwartet 'Begriff: Übersetzung')
        lines = ai_raw_text.split("\n")
        for line in lines:
            if ":" in line:
                parts = line.split(":", 1)
                orig = parts[0].strip()
                trans = parts[1].strip()

                # Nur hinzufügen, wenn der Begriff (kleingeschrieben) noch nicht existiert
                if orig.lower() not in existing_dictionary and orig:
                    existing_dictionary[orig.lower()] = trans
                    new_entries_count += 1

        log.info(f"Dictionary-Merge: {new_entries_count} neue Begriffe gefunden.")
        return existing_dictionary

    def load_dictionary_file(self):
        """Lädt das Dictionary aus einer Textdatei und gibt es zurück."""
        dict = {}
        if os.path.exists("dictionary.txt"):
            with open("dictionary.txt", "r", encoding="utf-8") as f:
                for line in f:
                    if ":" in line:
                        orig, trans = line.split(":", 1)
                        dict[orig.strip().lower()] = trans.strip()
        return dict

    def save_dictionary_file(self, glossary_dict):
        """Speichert das Dictionary sortiert in die Textdatei."""
        with open("dictionary.txt", "w", encoding="utf-8") as f:
            # Sortiert nach Originalbegriff
            for orig in sorted(glossary_dict.keys()):
                f.write(f"{orig.capitalize()}: {glossary_dict[orig]}\n")


    def save_chat(self):
        # 1. Gesamten Text aus UI holen
        raw_content = self.chat_area.get("1.0", tk.END)

        # 2. Nur die Übersetzungen extrahieren
        clean_content = self.extract_translations_only(raw_content)

        if not clean_content:
            messagebox.showwarning("Fehler", "Keine übersetzten Abschnitte gefunden. "
                                             "Achten Sie auf die Markierungen BEGINN/ENDE.")
            return

        # 3. Dateidialog
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[
                ("Text", "*.txt"),
                ("Word", "*.docx"),
                ("PDF", "*.pdf"),
                ("EPUB", "*.epub"),
                ("Markdown", "*.md")
            ]
        )

        if not file_path:
            return

        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == ".docx":
                from docx import Document
                doc = Document()
                # Text in Absätze teilen (nach echten Umbrüchen im Quelltext)
                for line in clean_content.split('\n'):
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)
                    else:
                        doc.add_paragraph("")  # Leerzeile erhalten
                doc.save(file_path)

            elif ext == ".pdf":
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=11)
                # UTF-8 zu Latin-1 Konvertierung mit Fallback für Sonderzeichen
                safe_text = clean_content.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 10, safe_text)
                pdf.output(file_path)

            elif ext == ".epub":
                from ebooklib import epub
                book = epub.EpubBook()
                basisname = os.path.basename(file_path).replace(ext, "")
                book.set_identifier(f"id_{int(time.time())}")
                book.set_title(basisname)
                book.set_language('de')

                # HTML-Inhalt erstellen (Zeilenumbrüche erhalten)
                html_body = clean_content.replace('\n', '<br/>')
                c1 = epub.EpubHtml(title='Übersetzung', file_name='chap_1.xhtml', lang='de')
                c1.content = f"<html><body><h1>{basisname}</h1><p>{html_body}</p></body></html>"

                book.add_item(c1)
                book.spine = ['nav', c1]
                epub.write_epub(file_path, book)

            else:  # TXT oder Markdown
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(clean_content)

            messagebox.showinfo("Erfolg", f"Datei wurde als {ext.upper()} gespeichert!")

        except Exception as e:
            log.error(f"Speicherfehler: {e}")
            messagebox.showerror("Fehler", f"Speichern fehlgeschlagen: {e}")

    def extract_translations_only(self, text):
        pattern = r"--- BEGINN ABSCHNITT \d+\n(.*?)\n--- ENDE ABSCHNITT \d+"
        matches = re.findall(pattern, text, re.DOTALL)

        # Falls Treffer gefunden wurden, geben wir sie (z.B. zusammengefügt) zurück
        if matches:
            return "\n".join(matches)

        # Falls kein Pattern gefunden wurde, aber der Text nicht leer ist:
        return text if text else ""

if __name__ == "__main__":
    root = tk.Tk()
    app = ChatApp(root)
    root.geometry("700x600")
    root.mainloop()